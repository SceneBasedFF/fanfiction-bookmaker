from __future__ import annotations

from dataclasses import dataclass, field
from html import escape
from pathlib import Path
from typing import Any
import re

import markdown as mdlib
import yaml
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


@dataclass(slots=True)
class Scene:
    name: str
    file: str
    text: str       # comments stripped; used for word counting
    raw_text: str  # original source text; preserved in .md output


@dataclass(slots=True)
class Chapter:
    slug: str
    name: str
    scenes: list[Scene] = field(default_factory=list)


@dataclass(slots=True)
class StoryConfig:
    title: str = "Untitled Book"
    author: str = ""
    description: str = ""
    subtitle: str = ""
    language: str = "en-GB"


@dataclass(slots=True)
class CompileResult:
    generated_files: list[Path]


CSS_TEXT = """
:root {
  color-scheme: light;
  --bg: #f5efe6;
  --panel: rgba(255, 255, 255, 0.72);
  --ink: #1f1a17;
  --muted: #675f57;
  --accent: #7a4c25;
  --line: rgba(122, 76, 37, 0.18);
}

* { box-sizing: border-box; }

html, body {
  margin: 0;
  padding: 0;
}

body {
  min-height: 100vh;
  background:
    radial-gradient(circle at top, rgba(122, 76, 37, 0.08), transparent 42%),
    linear-gradient(180deg, #fbf7f1 0%, var(--bg) 100%);
  color: var(--ink);
  font-family: "Iowan Old Style", "Palatino Linotype", Palatino, Georgia, serif;
  line-height: 1.75;
}

main {
  width: min(72ch, calc(100vw - 2rem));
  margin: 0 auto;
  padding: 4rem 0 6rem;
}

article {
  background: var(--panel);
  backdrop-filter: blur(6px);
  border: 1px solid rgba(255, 255, 255, 0.7);
  border-radius: 20px;
  padding: clamp(1.5rem, 4vw, 3rem);
  box-shadow: 0 16px 50px rgba(47, 35, 23, 0.08);
}

h1, h2, h3, h4, h5, h6 {
  font-family: "Aptos", "Segoe UI", system-ui, sans-serif;
  letter-spacing: 0.02em;
  line-height: 1.25;
  color: var(--ink);
  margin: 2rem 0 1rem;
}

h1 {
  margin-top: 0;
  font-size: clamp(2.2rem, 5vw, 3.6rem);
}

h2 {
  font-size: clamp(1.5rem, 3vw, 2rem);
  padding-top: 1rem;
  border-top: 1px solid var(--line);
}

h3 {
  font-size: 1.25rem;
}

p, ul, ol, blockquote, pre, table {
  margin: 0 0 1.1rem;
}

p {
  font-size: 1.05rem;
}

blockquote {
  border-left: 4px solid var(--accent);
  padding-left: 1rem;
  color: var(--muted);
  font-style: italic;
}

code {
  font-family: "Cascadia Mono", "SFMono-Regular", Consolas, monospace;
  font-size: 0.94em;
  background: rgba(122, 76, 37, 0.08);
  padding: 0.12em 0.35em;
  border-radius: 0.35rem;
}

pre {
  overflow-x: auto;
  padding: 1rem 1.1rem;
  border-radius: 14px;
  background: rgba(31, 26, 23, 0.06);
  border: 1px solid rgba(31, 26, 23, 0.08);
}

pre code {
  background: transparent;
  padding: 0;
}

hr {
  border: 0;
  border-top: 1px solid var(--line);
  margin: 2rem 0;
}

a {
  color: var(--accent);
  text-decoration-thickness: 0.08em;
  text-underline-offset: 0.12em;
}

img {
  max-width: 100%;
  height: auto;
}

table {
  width: 100%;
  border-collapse: collapse;
  font-size: 0.98rem;
}

th, td {
  border-bottom: 1px solid var(--line);
  text-align: left;
  padding: 0.45rem 0.5rem;
  vertical-align: top;
}
""".strip()


SAFE_NAME_RE = re.compile(r"^[A-Za-z0-9][A-Za-z0-9_-]*$")
HTML_COMMENT_RE = re.compile(r"<!--.*?-->", re.DOTALL)
MARKDOWN_COMMENT_LINE_RE = re.compile(r"^\s*\[//\]:\s*#\s*\(.*\)\s*$")
PREFIX_CHAIN_RE = re.compile(r"^[A-Za-z]+\d+(?:_[A-Za-z]+\d+)*$")
WORD_RE = re.compile(r"[^\W_]+(?:['’-][^\W_]+)*", re.UNICODE)


def compile_project(root: Path, config_filename: str, output_dir: str) -> CompileResult:
    root = root.resolve()
    config_path = safe_resolve_within(root, config_filename)
    config = load_config(config_path)
    chapter_order = load_chapter_order(root / "chapters.yml")
    chapters = [load_chapter(root, slug) for slug in chapter_order]

    output_root = safe_resolve_within(root, output_dir)
    chapter_output_root = output_root / "chapters"
    assets_output_root = output_root / "assets"
    chapter_output_root.mkdir(parents=True, exist_ok=True)
    assets_output_root.mkdir(parents=True, exist_ok=True)

    # Remove stale generated chapter files so renames do not leave old artifacts behind.
    for stale_path in chapter_output_root.glob("*.html"):
        stale_path.unlink(missing_ok=True)
    for stale_path in chapter_output_root.glob("*.docx"):
        stale_path.unlink(missing_ok=True)

    # Remove legacy top-level outputs from older versions.
    for stale_name in ("book.md", "book.html", "book.docx", "book-stats.md", "spellcheck-report.md", "spellcheck-report.json"):
        (output_root / stale_name).unlink(missing_ok=True)

    css_path = assets_output_root / "style.css"
    css_path.write_text(CSS_TEXT + "\n", encoding="utf-8")

    generated_files: list[Path] = [css_path]

    for chapter_number, chapter in enumerate(chapters, start=1):
        chapter_markdown = compose_chapter_markdown(chapter, strip_comments=True)
        chapter_title = format_chapter_title(chapter_number, chapter.name)
        chapter_slug = f"ch{chapter_number:02d}-{slugify(chapter.name, default=slugify(chapter.slug))}"
        chapter_html_path = chapter_output_root / f"{chapter_slug}.html"
        chapter_docx_path = chapter_output_root / f"{chapter_slug}.docx"
        write_html_document(
            markdown_text=chapter_markdown,
            title=chapter_title,
            output_path=chapter_html_path,
            css_href="../assets/style.css",
            language=config.language,
            subtitle=config.subtitle,
            author=config.author,
        )
        write_docx_document(
            markdown_text=chapter_markdown,
            output_path=chapter_docx_path,
            title=chapter_title,
            author=config.author,
            subtitle=config.subtitle,
            language=config.language,
        )
        generated_files.extend([chapter_html_path, chapter_docx_path])

    book_markdown = compose_book_markdown(config, chapters, strip_comments=False)
    book_markdown_stripped = compose_book_markdown(config, chapters, strip_comments=True)
    book_slug = slugify(config.title, default="book")
    root_readme_path = root / "README.md"
    book_md_path = output_root / f"{book_slug}.md"
    book_md_path.write_text(book_markdown, encoding="utf-8")
    root_readme_path.write_text(render_story_readme(config, book_md_path.relative_to(root)), encoding="utf-8")
    stats_md_path = output_root / f"{book_slug}-stats.md"
    stats_md_path.write_text(render_story_stats_markdown(config, chapters), encoding="utf-8")
    book_html_path = output_root / f"{book_slug}.html"
    book_docx_path = output_root / f"{book_slug}.docx"
    write_html_document(
        markdown_text=book_markdown_stripped,
        title=config.title,
        output_path=book_html_path,
        css_href="assets/style.css",
        language=config.language,
        subtitle=config.subtitle,
        author=config.author,
    )
    write_docx_document(
        markdown_text=book_markdown_stripped,
        output_path=book_docx_path,
        title=config.title,
        author=config.author,
        subtitle=config.subtitle,
        language=config.language,
    )
    generated_files.extend([root_readme_path, book_md_path, stats_md_path, book_html_path, book_docx_path])

    return CompileResult(generated_files=generated_files)


def load_config(path: Path) -> StoryConfig:
    if not path.exists():
        return StoryConfig()

    data = read_yaml(path)
    if data is None:
        return StoryConfig()
    if not isinstance(data, dict):
        raise ValueError(f"Config file {path} must contain a mapping.")

    return StoryConfig(
        title=str(data.get("title") or "Untitled Book"),
        author=str(data.get("author") or ""),
        description=str(data.get("description") or data.get("subtitle") or ""),
        subtitle=str(data.get("subtitle") or ""),
        language=str(data.get("language") or "en-GB"),
    )


def load_chapter_order(path: Path) -> list[str]:
    data = read_yaml(path)
    if not isinstance(data, dict):
        raise ValueError(f"{path} must contain a mapping with a chapters list.")

    chapter_list = data.get("chapters")
    if not isinstance(chapter_list, list) or not chapter_list:
        raise ValueError(f"{path} must contain a non-empty chapters list.")

    order: list[str] = []
    seen: set[str] = set()
    for item in chapter_list:
        if not isinstance(item, str) or not item.strip():
            raise ValueError(f"{path} contains an invalid chapter entry: {item!r}")
        slug = validate_safe_name(item.strip(), f"chapter slug in {path}")
        if slug in seen:
            raise ValueError(f"{path} contains a duplicate chapter entry: {slug}")
        seen.add(slug)
        order.append(slug)
    return order


def load_chapter(root: Path, slug: str) -> Chapter:
    safe_slug = validate_safe_name(slug, "chapter slug")
    path = resolve_referenced_file(root / "chapters", safe_slug, "yml")
    data = read_yaml(path)
    if not isinstance(data, dict):
        raise ValueError(f"{path} must contain a chapter mapping.")

    chapter_data = data.get("chapter")
    if not isinstance(chapter_data, dict):
        raise ValueError(f"{path} must contain a top-level chapter mapping.")

    name = str(chapter_data.get("name") or "").strip()
    if not name:
        raise ValueError(f"{path} is missing chapter.name.")

    scenes_data = chapter_data.get("scenes")
    if not isinstance(scenes_data, list) or not scenes_data:
        raise ValueError(f"{path} must contain a non-empty chapter.scenes list.")

    scenes: list[Scene] = []
    seen_files: set[str] = set()
    for index, scene_entry in enumerate(scenes_data, start=1):
        if not isinstance(scene_entry, dict):
            raise ValueError(f"{path} scene #{index} must be a mapping.")
        scene_name = str(scene_entry.get("name") or "").strip()
        scene_file = validate_safe_name(str(scene_entry.get("file") or "").strip(), f"scene file in {path}")
        if not scene_name:
            raise ValueError(f"{path} scene #{index} is missing a name.")
        if scene_file in seen_files:
            raise ValueError(f"{path} contains a duplicate scene file: {scene_file}")
        seen_files.add(scene_file)

        scene_path = resolve_referenced_file(root / "scenes", scene_file, "md")
        if not scene_path.exists():
            raise FileNotFoundError(f"Missing scene file: {scene_path}")

        raw = scene_path.read_text(encoding="utf-8")
        scenes.append(
            Scene(
                name=scene_name,
                file=scene_file,
                text=strip_scene_comments(raw),
                raw_text=raw,
            )
        )

    return Chapter(slug=slug, name=name, scenes=scenes)


def compose_chapter_markdown(chapter: Chapter, strip_comments: bool = False) -> str:
    lines: list[str] = []
    for scene in chapter.scenes:
        lines.append("<br>")
        lines.append(f"<h4 align=\"center\">{escape(scene.name)}</h4>")
        lines.append("<br>")
        body = strip_scene_comments(scene.raw_text) if strip_comments else scene.raw_text
        if body.strip():
            lines.append(body.strip())
    return "\n\n".join(lines).strip() + "\n"


def compose_book_markdown(config: StoryConfig, chapters: list[Chapter], strip_comments: bool = False) -> str:
    lines: list[str] = []

    for chapter_number, chapter in enumerate(chapters, start=1):
        lines.append(f"## {format_chapter_title(chapter_number, chapter.name)}")
        for scene in chapter.scenes:
            lines.append("<br>")
            lines.append(f"<h4 align=\"center\">{escape(scene.name)}</h4>")
            lines.append("<br>")
            body = strip_scene_comments(scene.raw_text) if strip_comments else scene.raw_text
            if body.strip():
                lines.append(body.strip())
    return "\n\n".join(lines).strip() + "\n"


def render_story_readme(config: StoryConfig, story_markdown_relative_path: Path) -> str:
    lines = [f"# {config.title}"]

    if config.author:
        lines.append(f"**Author:** {config.author}")

    if config.description:
        lines.append(config.description)

    lines.append(f"[Read the story]({story_markdown_relative_path.as_posix()})")
    return "\n\n".join(lines).strip() + "\n"


def render_story_stats_markdown(config: StoryConfig, chapters: list[Chapter]) -> str:
    lines = [f"# {config.title} Statistics", ""]

    if config.author:
        lines.append(f"**Author:** {config.author}")
        lines.append("")

    chapter_stats: list[tuple[int, Chapter, int, list[tuple[Scene, int]]]] = []
    total_words = 0

    for chapter_number, chapter in enumerate(chapters, start=1):
        scene_stats: list[tuple[Scene, int]] = []
        chapter_words = 0
        for scene in chapter.scenes:
            scene_words = count_words_in_markdown(scene.text)
            scene_stats.append((scene, scene_words))
            chapter_words += scene_words
        chapter_stats.append((chapter_number, chapter, chapter_words, scene_stats))
        total_words += chapter_words

    lines.extend(
        [
            "## Overview",
            "",
            f"- Total word count: {total_words:,}",
            f"- Total chapters: {len(chapters)}",
            f"- Total scenes: {sum(len(chapter.scenes) for chapter in chapters)}",
            "",
            "## Chapter Breakdown",
            "",
        ]
    )

    for chapter_number, chapter, chapter_words, scene_stats in chapter_stats:
        lines.append(f"### {format_chapter_title(chapter_number, chapter.name)}")
        lines.append("")
        lines.append(f"- Chapter word count: {chapter_words:,}")
        lines.append(f"- Scene count: {len(scene_stats)}")
        lines.append("")
        lines.append("| Scene | Words |")
        lines.append("| --- | ---: |")
        for scene, scene_words in scene_stats:
            lines.append(f"| {escape_pipe(scene.name)} | {scene_words:,} |")
        lines.append("")

    return "\n".join(lines).rstrip() + "\n"


def format_chapter_title(chapter_number: int, chapter_name: str) -> str:
    return f"Chapter {chapter_number}: {chapter_name}"


def write_html_document(
    markdown_text: str,
    title: str,
    output_path: Path,
    css_href: str,
    language: str,
    subtitle: str = "",
    author: str = "",
) -> None:
    body_html = markdown_to_html(markdown_text)
    document = f"""<!doctype html>
<html lang="{escape(language)}">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <meta name="generator" content="fanfic-bookmaker">
  <title>{escape(title)}</title>
  <link rel="stylesheet" href="{escape(css_href)}">
</head>
<body>
  <main>
    <article>
      {render_front_matter(title, author, subtitle)}
      {body_html}
    </article>
  </main>
</body>
</html>
"""
    output_path.write_text(document, encoding="utf-8")


def render_front_matter(title: str, author: str, subtitle: str) -> str:
    parts = [f"<h1>{escape(title)}</h1>"]
    if subtitle:
        parts.append(f"<p><em>{escape(subtitle)}</em></p>")
    if author:
        parts.append(f"<p><strong>{escape(author)}</strong></p>")
    return "\n      ".join(parts)


def markdown_to_html(markdown_text: str) -> str:
    return mdlib.markdown(
        markdown_text,
        extensions=[
            "extra",
            "sane_lists",
            "fenced_code",
            "tables",
            "attr_list",
            "md_in_html",
        ],
        output_format="html5",
    )


def write_docx_document(
    markdown_text: str,
    output_path: Path,
    title: str,
    author: str,
    subtitle: str,
    language: str,
) -> None:
    html = markdown_to_html(markdown_text)
    soup = BeautifulSoup(html, "html.parser")
    document = Document()

    document.core_properties.title = title
    if author:
        document.core_properties.author = author
    document.core_properties.language = language
    document.styles["Normal"].font.name = "Georgia"
    document.styles["Normal"].font.size = Pt(11)

    title_paragraph = document.add_paragraph(style="Title")
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(title)
    title_run.bold = True

    if subtitle:
        subtitle_paragraph = document.add_paragraph()
        subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_paragraph.add_run(subtitle)
        subtitle_run.italic = True

    if author:
        author_paragraph = document.add_paragraph()
        author_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author_paragraph.add_run(author)
        author_run.bold = True

    section = document.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    for node in soup.contents:
        append_block(document, node)

    document.save(output_path)


def append_block(document: Document, node: Any, list_depth: int = 0) -> None:
    if isinstance(node, NavigableString):
        if str(node).strip():
            paragraph = document.add_paragraph()
            paragraph.add_run(str(node).strip())
        return

    if not isinstance(node, Tag):
        return

    name = node.name.lower()
    if name in {"html", "body", "main", "article"}:
        for child in node.children:
            append_block(document, child, list_depth=list_depth)
        return

    if name in {"h1", "h2", "h3", "h4", "h5", "h6"}:
        level = int(name[1])
        heading_level = max(level - 1, 1)
        paragraph = document.add_paragraph(style=f"Heading {min(heading_level, 9)}")
        append_inline(paragraph, node)
        return

    if name == "p":
        paragraph = document.add_paragraph()
        append_inline(paragraph, node)
        return

    if name in {"ul", "ol"}:
        for li in node.find_all("li", recursive=False):
            render_list_item(document, li, ordered=name == "ol", depth=list_depth)
        return

    if name == "blockquote":
        paragraph = document.add_paragraph(style="Intense Quote")
        for child in node.children:
            append_inline(paragraph, child)
        return

    if name == "pre":
        paragraph = document.add_paragraph()
        paragraph.style = "No Spacing"
        run = paragraph.add_run(node.get_text())
        run.font.name = "Cascadia Mono"
        run.font.size = Pt(9.5)
        return

    if name == "hr":
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run("-" * 18)
        return

    if name == "table":
        render_table(document, node)
        return

    for child in node.children:
        append_block(document, child, list_depth=list_depth)


def render_list_item(document: Document, li: Tag, ordered: bool, depth: int) -> None:
    style = "List Number" if ordered else "List Bullet"
    paragraph = document.add_paragraph(style=style)
    paragraph.paragraph_format.left_indent = Inches(0.25 * depth)

    inline_children = []
    for child in li.children:
        if isinstance(child, Tag) and child.name in {"ul", "ol"}:
            continue
        inline_children.append(child)

    for child in inline_children:
        append_inline(paragraph, child)

    for child in li.find_all(["ul", "ol"], recursive=False):
        append_block(document, child, list_depth=depth + 1)


def render_table(document: Document, table_tag: Tag) -> None:
    rows = table_tag.find_all("tr", recursive=False)
    if not rows:
        return
    max_cols = max(len(row.find_all(["th", "td"], recursive=False)) for row in rows)
    table = document.add_table(rows=len(rows), cols=max(max_cols, 1))
    table.style = "Table Grid"

    for row_idx, row_tag in enumerate(rows):
        cells = row_tag.find_all(["th", "td"], recursive=False)
        for col_idx, cell_tag in enumerate(cells):
            if col_idx >= len(table.rows[row_idx].cells):
                continue
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = normalize_whitespace(cell_tag.get_text(" ", strip=True))


def append_inline(paragraph, node: Any, bold: bool = False, italic: bool = False, code: bool = False) -> None:
    if isinstance(node, NavigableString):
        text = str(node)
        if not text:
            return
        run = paragraph.add_run(text)
        apply_run_style(run, bold=bold, italic=italic, code=code)
        return

    if not isinstance(node, Tag):
        return

    name = node.name.lower()
    if name in {"strong", "b"}:
        for child in node.children:
            append_inline(paragraph, child, bold=True, italic=italic, code=code)
        return
    if name in {"em", "i"}:
        for child in node.children:
            append_inline(paragraph, child, bold=bold, italic=True, code=code)
        return
    if name == "code":
        for child in node.children:
            append_inline(paragraph, child, bold=bold, italic=italic, code=True)
        return
    if name == "a":
        add_hyperlink(paragraph, node.get_text(), node.get("href", ""), bold=bold, italic=italic)
        return
    if name == "br":
        paragraph.add_run().add_break()
        return

    for child in node.children:
        append_inline(paragraph, child, bold=bold, italic=italic, code=code)


def apply_run_style(run, bold: bool, italic: bool, code: bool) -> None:
    run.bold = bold
    run.italic = italic
    if code:
        run.font.name = "Cascadia Mono"
        run.font.size = Pt(9.5)


def add_hyperlink(paragraph, text: str, url: str, bold: bool = False, italic: bool = False):
    if not url:
        run = paragraph.add_run(text)
        apply_run_style(run, bold=bold, italic=italic, code=False)
        return

    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    if bold:
        b = OxmlElement("w:b")
        r_pr.append(b)
    if italic:
        i = OxmlElement("w:i")
        r_pr.append(i)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "7A4C25")
    r_pr.append(color)
    under = OxmlElement("w:u")
    under.set(qn("w:val"), "single")
    r_pr.append(under)
    new_run.append(r_pr)

    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def normalize_whitespace(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def count_words_in_markdown(text: str) -> int:
    plain_text = markdown_to_plain_text(text)
    return len(WORD_RE.findall(plain_text))


def markdown_to_plain_text(text: str) -> str:
    html = markdown_to_html(text)
    soup = BeautifulSoup(html, "html.parser")
    return soup.get_text(" ", strip=True)


def strip_scene_comments(text: str) -> str:
    without_html_comments = HTML_COMMENT_RE.sub("", text)
    lines = without_html_comments.splitlines()
    filtered_lines = [line for line in lines if not MARKDOWN_COMMENT_LINE_RE.match(line)]
    return "\n".join(filtered_lines)


def escape_pipe(text: str) -> str:
    return text.replace("|", "\\|")


def resolve_referenced_file(base_dir: Path, reference_name: str, extension: str) -> Path:
    exact_path = safe_resolve_within(base_dir, f"{reference_name}.{extension}")
    if exact_path.exists():
        return exact_path

    pattern = f"*_{reference_name}.{extension}"
    matches: list[Path] = []
    for candidate in base_dir.glob(pattern):
        stem = candidate.stem
        prefix = stem[: -(len(reference_name) + 1)]
        if PREFIX_CHAIN_RE.fullmatch(prefix):
            matches.append(candidate.resolve())

    if not matches:
        return exact_path

    unique_matches = sorted(set(matches))
    if len(unique_matches) > 1:
        joined = ", ".join(str(path) for path in unique_matches)
        raise ValueError(f"Ambiguous file reference '{reference_name}.{extension}' in {base_dir}: {joined}")

    match = unique_matches[0]
    if match != base_dir.resolve() and not match.is_relative_to(base_dir.resolve()):
        raise ValueError(f"Resolved file escapes project directory: {match}")
    return match


def read_yaml(path: Path) -> Any:
    if not path.exists():
        raise FileNotFoundError(f"Missing required file: {path}")
    with path.open("r", encoding="utf-8") as handle:
        return yaml.safe_load(handle)


def validate_safe_name(value: str, field_name: str) -> str:
    if not value:
        raise ValueError(f"{field_name} cannot be empty.")
    if not SAFE_NAME_RE.fullmatch(value):
        raise ValueError(
            f"{field_name} contains unsupported characters. "
            "Use only letters, numbers, hyphens, and underscores."
        )
    return value


def safe_resolve_within(base_dir: Path, relative_or_absolute: str) -> Path:
    base_resolved = base_dir.resolve()
    candidate = (base_resolved / relative_or_absolute).resolve()
    if candidate != base_resolved and not candidate.is_relative_to(base_resolved):
        raise ValueError(f"Path escapes project directory: {relative_or_absolute}")
    return candidate


def slugify(value: str, default: str = "chapter") -> str:
    value = value.strip().lower()
    value = re.sub(r"[^a-z0-9]+", "-", value)
    return value.strip("-") or default
